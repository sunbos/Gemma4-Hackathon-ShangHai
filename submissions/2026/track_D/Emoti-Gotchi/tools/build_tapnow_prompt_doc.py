from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Emoti-Gotchi_TapNow儿童情景视频制作手册.docx"
ASSETS = ROOT / "src" / "assets"

INK = RGBColor(61, 45, 37)
MUTED = RGBColor(112, 92, 78)
ORANGE = RGBColor(225, 118, 65)
GREEN = RGBColor(82, 151, 122)
LIGHT = "F8F1E8"
PALE_GREEN = "EDF6F1"
PALE_ORANGE = "FFF0E4"

def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=140, start=160, bottom=140, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def style_run(run, size=10.5, bold=False, color=INK, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color

def add_para(doc, text="", size=10.5, bold=False, color=INK, before=0, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.28
    if align:
        p.alignment = align
    style_run(p.add_run(text), size=size, bold=bold, color=color)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    style_run(p.add_run(text), size=10.5, color=INK)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    style_run(run, size=16 if level == 1 else 12.5, bold=True, color=ORANGE if level == 1 else INK)
    return p

def add_callout(doc, label, body, fill=PALE_GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.25)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run(label), size=10, bold=True, color=GREEN if fill == PALE_GREEN else ORANGE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.3
    p2.paragraph_format.space_after = Pt(0)
    style_run(p2.add_run(body), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_prompt_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.25)
    set_cell_fill(cell, LIGHT)
    set_cell_margins(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(0)
    style_run(p.add_run(text), size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_image_pair(doc, paths):
    table = doc.add_table(rows=1, cols=len(paths))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, path in enumerate(paths):
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, 50, 50, 50, 50)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(2.55 if len(paths) == 2 else 3.7))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK

for name, size, color in (("Heading 1", 16, ORANGE), ("Heading 2", 12.5, INK), ("Heading 3", 11, GREEN)):
    style = doc.styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
style_run(header.add_run("EMOTI-GOTCHI · TAPNOW VIDEO GUIDE"), size=8.5, bold=True, color=MUTED)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(footer.add_run("儿童隐私优先的端侧情绪支持系统 · AI for Social Good"), size=8, color=MUTED)

add_para(doc, "EMOTI-GOTCHI", size=11, bold=True, color=ORANGE, after=8)
add_para(doc, "TapNow 儿童使用情景\n视频制作手册", size=27, bold=True, color=INK, after=8)
add_para(doc, "用于产品形象片、赛事演示视频与儿童家庭使用场景展示", size=12, color=MUTED, after=18)
add_image_pair(doc, [ASSETS / "emoti-gotchi-happy.png"])
add_callout(
    doc,
    "核心叙事",
    "孩子需要的，有时不是答案，而是一个愿意先陪他安静下来的伙伴。视频重点展示儿童自然使用情景，而不是抽象科技能力。",
    PALE_ORANGE,
)

add_heading(doc, "1. 制作目标", 1)
add_bullet(doc, "展示 6–8 岁儿童在日常家庭情境中自然使用 Emoti-Gotchi。")
add_bullet(doc, "表现玩偶如何先陪伴、帮助表达，并让家长获得温和且不侵犯隐私的行动提示。")
add_bullet(doc, "明确产品不是监控工具、诊断工具或专业心理支持的替代品。")
add_bullet(doc, "在高风险情况下体现端侧安全机制、家长授权与专业支持转介。")

add_heading(doc, "2. TapNow 统一生成设置", 1)
settings = doc.add_table(rows=1, cols=2)
settings.alignment = WD_TABLE_ALIGNMENT.CENTER
settings.autofit = False
settings.columns[0].width = Inches(1.65)
settings.columns[1].width = Inches(4.6)
for c, text in zip(settings.rows[0].cells, ["设置项", "推荐值"]):
    set_cell_fill(c, PALE_ORANGE)
    set_cell_margins(c)
    style_run(c.paragraphs[0].add_run(text), size=10, bold=True, color=INK)
for left, right in [
    ("画面比例", "16:9"),
    ("单镜头时长", "5–7 秒"),
    ("运动强度", "低；避免玩偶与儿童动作夸张"),
    ("镜头运动", "固定镜头或缓慢推进"),
    ("角色一致性", "参考图权重尽可能调高"),
    ("字幕", "TapNow 内不生成，后期统一添加"),
]:
    row = settings.add_row()
    for cell in row.cells:
        set_cell_margins(cell)
    style_run(row.cells[0].paragraphs[0].add_run(left), size=10, bold=True, color=INK)
    style_run(row.cells[1].paragraphs[0].add_run(right), size=10, color=INK)

add_heading(doc, "3. 每个镜头都需附加的约束词", 1)
add_heading(doc, "统一角色与儿童隐私约束", 2)
add_prompt_box(doc, "保持参考图中 Emoti-Gotchi 玩偶的外形、比例、材质与圆形屏幕脸完全一致。孩子约 7 岁，仅展示侧脸、背影或手部，避免清晰正脸。真实温暖的家庭生活场景，自然儿童动作，玩偶动作轻柔克制，不说教、不替代父母。")
add_heading(doc, "统一负面提示词", 2)
add_prompt_box(doc, "儿童正脸特写，儿童身份信息，悲惨哭泣，恐怖气氛，角色变形，多余肢体，新增腿部，玩偶主动监控儿童，摄像头，文字，水印，剧烈动作，夸张卡通动画，背景扭曲，手指变形")

scenes = [
    {
        "title": "镜头 1｜放学后的分享",
        "image": [ASSETS / "emoti-gotchi-happy.png"],
        "purpose": "建立产品第一印象：玩偶是孩子愿意自然分享日常的小伙伴。",
        "prompt": "一名约 7 岁的孩子放学回家，将书包放下，坐在卧室地毯上，开心地向 Emoti-Gotchi 分享今天和朋友玩游戏的事情。孩子只展示背影和手部。玩偶被孩子抱在怀里，屏幕脸露出温暖笑容，轻轻发出暖黄色光，像在认真倾听。真实自然的儿童家庭生活，缓慢推进镜头。",
        "caption": "每一种小事，都值得被认真听见。",
    },
    {
        "title": "镜头 2｜玩具损坏后的生气",
        "image": [ASSETS / "emoti-gotchi-angry.png", ASSETS / "emoti-gotchi-calm.png"],
        "purpose": "证明玩偶不会跟随发怒，而是验证情绪并帮助共同调节。",
        "prompt": "孩子坐在地毯上，因为积木作品倒塌而生气，将手臂抱在胸前。Emoti-Gotchi 安静陪在孩子身边，不模仿愤怒，也不立即说教。玩偶屏幕先显示理解和关切，随后缓慢转为平静的绿色呼吸光，引导孩子一起慢慢平复。孩子肩膀逐渐放松。保持低刺激、温和陪伴。",
        "caption": "先接住情绪，再慢慢平静。",
    },
    {
        "title": "镜头 3｜睡前害怕黑暗",
        "image": [ASSETS / "emoti-gotchi-sad.png", ASSETS / "emoti-gotchi-calm.png"],
        "purpose": "表现儿童端即时支持不等待云端，玩偶通过低刺激方式陪伴。",
        "prompt": "夜晚，孩子躺在床上，因为房间较暗而有些害怕，轻轻抱住 Emoti-Gotchi。只展示孩子侧后方轮廓，不展示清晰正脸。玩偶屏幕显示温柔关切的表情，随后发出柔和绿色呼吸光。床头灯逐渐变得温暖稳定，孩子慢慢放松。场景安静、安全，不渲染恐惧。",
        "caption": "即时陪伴在端侧，断网也能安心。",
    },
    {
        "title": "镜头 4｜“我没事”的隐藏委屈",
        "image": [ASSETS / "emoti-gotchi-happy.png", ASSETS / "emoti-gotchi-sad.png"],
        "purpose": "体现多通道推理的价值：发现语言与状态之间的信号冲突。",
        "prompt": "孩子放学后安静坐在窗边，嘴上说自己没事，但低着头、握紧书包带。Emoti-Gotchi 放在孩子身边，认真倾听。玩偶的屏幕脸从普通微笑缓慢转为温柔关切的委屈表情，表示注意到了语言与状态之间的差异。玩偶不逼问，只安静靠近孩子。真实克制，体现理解而不是监控。",
        "caption": "听见语言之外的情绪。",
    },
    {
        "title": "镜头 5｜家长获得温和提示",
        "image": [ASSETS / "emoti-gotchi-calm.png"],
        "purpose": "把父母端呈现为家庭天气预报，而不是儿童行为监控。",
        "prompt": "孩子在房间里与 Emoti-Gotchi 安静阅读，画面切换到客厅中的家长。家长手机只显示抽象的家庭天气图标和一句简单环境建议，不显示孩子原话、录音或具体行为记录。家长轻轻调暗刺眼灯光，并安静地坐到孩子附近陪伴。表现产品帮助家长用共情替代追问。",
        "caption": "不是监控孩子，而是帮助家庭改善环境。",
    },
    {
        "title": "镜头 6｜安全情况与专业支持连接",
        "image": [ASSETS / "emoti-gotchi-sad.png"],
        "purpose": "展示安全兜底，同时保留家长决策权并避免制造恐慌。",
        "prompt": "孩子抱着 Emoti-Gotchi 安静坐着，玩偶保持柔和、稳定的关切表情，不展示危险内容。画面切换到家长手机收到克制的安全提醒。家长深呼吸后平静走向孩子，坐在孩子身旁。在家长明确授权后，手机显示正在连接专业支持资源。整个过程温和、安全，不制造恐慌。",
        "caption": "生命安全优先，专业支持由家长明确授权。",
    },
]

for i, scene in enumerate(scenes):
    doc.add_page_break()
    add_heading(doc, scene["title"], 1)
    add_image_pair(doc, scene["image"])
    add_callout(doc, "镜头目的", scene["purpose"], PALE_GREEN)
    add_heading(doc, "TapNow 图生视频提示词", 2)
    add_prompt_box(doc, scene["prompt"])
    add_heading(doc, "建议字幕", 2)
    add_callout(doc, "ON-SCREEN COPY", scene["caption"], PALE_ORANGE)

doc.add_page_break()
add_heading(doc, "4. 推荐成片结构", 1)
add_para(doc, "建议将六个镜头剪辑为一条 45–60 秒产品形象片，并放在赛事演示视频开头或结尾。", color=MUTED, after=10)
timeline = doc.add_table(rows=1, cols=3)
timeline.alignment = WD_TABLE_ALIGNMENT.CENTER
timeline.autofit = False
for cell, text in zip(timeline.rows[0].cells, ["时间", "镜头", "核心信息"]):
    set_cell_fill(cell, PALE_ORANGE)
    set_cell_margins(cell)
    style_run(cell.paragraphs[0].add_run(text), size=10, bold=True, color=INK)
for item in [
    ("0–8 秒", "放学后的分享", "孩子愿意自然表达"),
    ("8–17 秒", "生气与共同调节", "玩偶不模仿发怒、不说教"),
    ("17–26 秒", "睡前害怕", "端侧即时陪伴"),
    ("26–35 秒", "隐藏委屈", "多通道识别信号冲突"),
    ("35–45 秒", "家长天气反馈", "脱敏、非监控式家庭支持"),
    ("45–60 秒", "安全兜底与结尾", "家长授权连接专业支持"),
]:
    row = timeline.add_row()
    for cell in row.cells:
        set_cell_margins(cell)
    for cell, text in zip(row.cells, item):
        style_run(cell.paragraphs[0].add_run(text), size=9.5, bold=(cell == row.cells[0]), color=INK)

add_heading(doc, "5. 开头与结尾文案", 1)
add_callout(doc, "开头", "孩子需要的，有时不是答案。\n而是一个愿意先陪他安静下来的伙伴。", PALE_ORANGE)
add_callout(doc, "结尾", "Emoti-Gotchi\n让孩子愿意表达，让家长更懂得回应。", PALE_GREEN)

add_heading(doc, "6. 最终审核清单", 1)
for text in [
    "所有镜头中的玩偶外形、屏幕脸位置、材质与比例保持一致。",
    "不出现儿童清晰正脸、身份信息、原始录音或完整对话。",
    "儿童情绪表现真实克制，不使用悲惨、恐怖或煽情画面。",
    "玩偶始终表现为陪伴媒介，不替代家长或专业心理人员。",
    "家长端只展示家庭气候、环境建议与安全提醒，不展示儿童行为侧写。",
    "专业支持连接必须表现为家长明确授权后的流程。",
    "字幕避免使用诊断、治疗、准确识别或证明改善等表述。",
]:
    add_bullet(doc, text)

add_callout(
    doc,
    "产品边界声明",
    "Emoti-Gotchi 是非诊断式家庭情绪支持系统。演示内容使用模拟数据，不代表已完成真实儿童测试、心理专家验证或真实硬件部署。",
    PALE_ORANGE,
)

doc.save(OUT)
print(OUT)
